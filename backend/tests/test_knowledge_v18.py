from __future__ import annotations

from pathlib import Path

import pytest


@pytest.fixture
def sample_long_txt(tmp_path):
    p = tmp_path / "doc.txt"
    text = "\n\n".join([f"Đoạn văn số {i} về trí tuệ nhân tạo và học máy." for i in range(20)])
    p.write_bytes(text.encode("utf-8"))
    return p


@pytest.fixture
def sample_docx(tmp_path):
    p = tmp_path / "sample.docx"
    try:
        from docx import Document
        doc = Document()
        doc.add_heading("Test Document", 0)
        doc.add_paragraph("Nội dung test cho knowledge base.")
        doc.save(str(p))
    except ImportError:
        p.write_bytes(b"fake docx")
    return p


class TestVectorStore:
    def test_init_creates_empty(self, app_module):
        from app.services.vector_store import get_vector_store
        vs = get_vector_store()
        assert vs.get_count() == 0

    def test_add_and_search_vectors(self, app_module):
        from app.services.vector_store import get_vector_store, VectorStore
        import importlib
        import sys
        store_module = sys.modules.get("app.services.vector_store")
        if store_module:
            importlib.reload(store_module)
        store_module.get_vector_store.cache_clear()
        vs = get_vector_store()
        vs.clear()
        texts = [
            "Trí tuệ nhân tạo đang phát triển nhanh chóng.",
            "Học máy là một nhánh của trí tuệ nhân tạo.",
            "Thời tiết hôm nay rất đẹp.",
        ]
        added = vs.add_vectors([1, 2, 3], texts)
        assert added == 3
        results = vs.search("trí tuệ nhân tạo", limit=2)
        assert len(results) >= 1
        assert results[0]["chunk_id"] in (1, 2)
        vs.clear()

    def test_remove_vectors(self, app_module):
        from app.services.vector_store import get_vector_store, VectorStore
        import importlib
        import sys
        store_module = sys.modules.get("app.services.vector_store")
        if store_module:
            importlib.reload(store_module)
        store_module.get_vector_store.cache_clear()
        vs = get_vector_store()
        vs.clear()
        vs.add_vectors([10, 20, 30], ["A B C", "D E F", "G H I"])
        assert vs.get_count() == 3
        removed = vs.remove_vectors({20})
        assert removed == 1
        assert vs.get_count() == 2
        vs.clear()
        assert vs.get_count() == 0


class TestHybridSearch:
    def test_hybrid_search_returns_results(self, app_module, sample_long_txt):
        from app.services.document_importer import import_document
        from app.services.knowledge_store import get_knowledge_store
        import_document(sample_long_txt)
        store = get_knowledge_store()
        results = store.hybrid_search_chunks("trí tuệ nhân tạo", limit=3, mode="hybrid")
        assert len(results) >= 1
        for r in results:
            assert "score" in r
        store.clear_knowledge_base()

    def test_keyword_mode_fallback(self, app_module, sample_long_txt):
        from app.services.document_importer import import_document
        from app.services.knowledge_store import get_knowledge_store
        import_document(sample_long_txt)
        store = get_knowledge_store()
        results = store.hybrid_search_chunks("trí tuệ", limit=3, mode="keyword")
        assert len(results) >= 1
        store.clear_knowledge_base()

    def test_semantic_mode_returns(self, app_module, sample_long_txt):
        from app.services.document_importer import import_document
        from app.services.knowledge_store import get_knowledge_store
        import_document(sample_long_txt)
        store = get_knowledge_store()
        results = store.hybrid_search_chunks("trí tuệ nhân tạo", limit=3, mode="semantic")
        assert len(results) >= 1
        store.clear_knowledge_base()

    def test_empty_query_returns_empty(self, app_module):
        from app.services.knowledge_store import get_knowledge_store
        store = get_knowledge_store()
        results = store.hybrid_search_chunks("", limit=5)
        assert results == []


class TestKnowledgeAPIv18:
    def test_search_with_mode_parameter(self, client, app_module, sample_long_txt):
        from app.services.document_importer import import_document
        import_document(sample_long_txt)
        response = client.post("/knowledge/search?query=trí tuệ&limit=3&mode=hybrid")
        assert response.status_code == 200
        data = response.json()
        assert "results" in data
        from app.services.knowledge_store import get_knowledge_store
        get_knowledge_store().clear_knowledge_base()

    def test_search_with_payload(self, client, app_module, sample_long_txt):
        from app.services.document_importer import import_document
        import_document(sample_long_txt)
        response = client.post(
            "/knowledge/search",
            json={"query": "trí tuệ", "limit": 3, "mode": "hybrid"},
        )
        assert response.status_code == 200
        data = response.json()
        assert "results" in data
        from app.services.knowledge_store import get_knowledge_store
        get_knowledge_store().clear_knowledge_base()

    def test_ask_with_mode(self, client, app_module, sample_long_txt):
        from app.services.document_importer import import_document
        import_document(sample_long_txt)
        response = client.post("/knowledge/ask?query=AI là gì?&mode=hybrid")
        assert response.status_code == 200
        data = response.json()
        assert "answer" in data
        assert "sources" in data
        from app.services.knowledge_store import get_knowledge_store
        get_knowledge_store().clear_knowledge_base()

    def test_document_preview(self, client, app_module, sample_long_txt):
        from app.services.document_importer import import_document
        result = import_document(sample_long_txt)
        doc_id = result["document_id"]
        response = client.get(f"/knowledge/documents/{doc_id}/content")
        assert response.status_code == 200
        data = response.json()
        assert "content" in data
        assert data["id"] == doc_id
        assert len(data["content"]) > 0
        from app.services.knowledge_store import get_knowledge_store
        get_knowledge_store().clear_knowledge_base()

    def test_document_preview_not_found(self, client, app_module):
        response = client.get("/knowledge/documents/99999/content")
        assert response.status_code == 404

    def test_export_knowledge(self, client, app_module, sample_long_txt):
        from app.services.document_importer import import_document
        import_document(sample_long_txt)
        response = client.get("/knowledge/export")
        assert response.status_code == 200
        data = response.json()
        assert "records" in data
        assert len(data["records"]) >= 1
        assert "document" in data["records"][0]
        assert "chunks" in data["records"][0]
        from app.services.knowledge_store import get_knowledge_store
        get_knowledge_store().clear_knowledge_base()

    def test_export_empty_knowledge(self, client, app_module):
        response = client.get("/knowledge/export")
        assert response.status_code == 200
        data = response.json()
        assert data["records"] == []

    def test_import_knowledge(self, client, app_module, sample_long_txt):
        from app.services.document_importer import import_document
        from app.services.knowledge_store import get_knowledge_store
        store = get_knowledge_store()
        import_document(sample_long_txt)
        export_resp = client.get("/knowledge/export")
        export_data = export_resp.json()
        store.clear_knowledge_base()
        assert len(store.list_documents()) == 0
        response = client.post("/knowledge/import", json={
            "records": export_data["records"],
            "mode": "merge",
        })
        assert response.status_code == 200
        data = response.json()
        assert data["documents_found"] >= 1
        assert data["documents_imported"] >= 1
        store.clear_knowledge_base()

    def test_import_duplicate_detected(self, client, app_module, sample_long_txt):
        from app.services.document_importer import import_document
        from app.services.knowledge_store import get_knowledge_store
        store = get_knowledge_store()
        import_document(sample_long_txt)
        export_resp = client.get("/knowledge/export")
        export_data = export_resp.json()
        response = client.post("/knowledge/import", json={
            "records": export_data["records"],
            "mode": "merge",
        })
        assert response.status_code == 200
        data = response.json()
        assert len(data["errors"]) > 0
        assert "duplicate" in data["errors"][0].lower()
        store.clear_knowledge_base()

    def test_import_replace_mode(self, client, app_module, sample_long_txt):
        from app.services.document_importer import import_document
        from app.services.knowledge_store import get_knowledge_store
        store = get_knowledge_store()
        import_document(sample_long_txt)
        export_resp = client.get("/knowledge/export")
        export_data = export_resp.json()
        response = client.post("/knowledge/import", json={
            "records": export_data["records"],
            "mode": "replace",
        })
        assert response.status_code == 200
        data = response.json()
        assert data["documents_imported"] >= 1
        store.clear_knowledge_base()


class TestCitations:
    def test_answer_includes_citations(self, app_module, sample_long_txt):
        from app.services.document_importer import import_document
        from app.services.knowledge_store import get_knowledge_store
        from app.services.knowledge_query import answer_from_knowledge
        import_document(sample_long_txt)
        import asyncio
        result = asyncio.run(answer_from_knowledge("trí tuệ nhân tạo là gì?", mode="hybrid"))
        assert "answer" in result
        assert "sources" in result
        if result["has_sufficient_context"]:
            for source in result["sources"]:
                assert "citation_index" in source
                assert source["citation_index"] >= 1
        get_knowledge_store().clear_knowledge_base()
